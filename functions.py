import os
import re
import pandas as pd
import openai
from docx import Document
import shutil
import boto3
from botocore.exceptions import NoCredentialsError
import webbrowser

ROOT_PATH = os.path.dirname(os.path.abspath(__file__))

# Utility functions
def create_folder_structure(base_path):
    subfolders = {
        "automation": ["automation script"],
        "developer": ["developer script"],
        "requirement": ["epic chunks", "epic response"],
        "tester": ["regression", "test case"]
    }
    for folder, subs in subfolders.items():
        folder_path = os.path.join(base_path, folder)
        os.makedirs(folder_path, exist_ok=True)
        for sub in subs:
            os.makedirs(os.path.join(folder_path, sub), exist_ok=True)

def generate_master_prompt(requirement_name, region, sections, specification):
    app_type = requirement_name  # Assuming app_type is the same as requirement_name
    if not region and not sections and not specification:
        master_prompt = (f"Act as a solution architect, your task is to create an end-to-end application for {app_type}. "
                         f"Your task is to generate a roadmap with just the Epics and the "
                         f"Features alone. Make sure that you provide only the Epics and Features for the SDLC of {app_type} "
                         f"use case and not any extra explanations, need not give any title to the response."
                         f"The template should be STRICTLY in below format, DO NOT INCLUDE '-' before Epic:\n"
                         "Epic 1: epic name\n"
                         "- Feature 1.1: feature name\n"
                         "- Feature 1.2: feature name\n")
    else:
        master_prompt = (f"Act as a solution architect, your task is to create an end-to-end application for {app_type} for {region} region. "
                         f"Make sure you include {sections}. Additional specifications are {specification}. Your task is to generate a roadmap with just the Epics and the "
                         f"Features alone. Make sure that you provide only the Epics and Features for the SDLC of {app_type} "
                         f"use case and not any extra explanations, need not give any title to the response."
                         f"The template should be STRICTLY in below format:\n"
                         "Epic 1: epic name\n"
                         "- Feature 1.1: feature name\n"
                         "- Feature 1.2: feature name\n")
    return master_prompt

def save_spec_prompt_csv(requirement_name, region, sections, specification, master_prompt):
    data = {
        "requirement_name": [requirement_name],
        "region": [region],
        "sections": [sections],
        "specification": [specification],
        "master_prompt": [master_prompt]
    }
    df = pd.DataFrame(data)
    csv_path = os.path.join(ROOT_PATH, "spec_prompt.csv")
    df.to_csv(csv_path, index=False)
    return csv_path

def fetch_openai_response(query):
    openai.api_key = 'sk-proj-Fff6MfBQrCqdmn5aqP0ZT3BlbkFJga08Ct6cUhRRCcX5qbqs'
    response = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",  # Use GPT-3.5 Turbo
    messages=[
        {"role":"system","content":"You are from JAVA background"},
        {"role":"user","content": f'{query}'}
    ],
    max_tokens=1024, 
    n=1,
    stop=None,
    temperature=0.1, 
    )
        
    return response['choices'][0]['message']['content']

def save_response_text_to_file(folder_path, requirement_name, response_text):
    # Define the file path
    requirement_folder = os.path.join(folder_path, "requirement")
    os.makedirs(requirement_folder, exist_ok=True)  # Ensure the folder exists
    file_path = os.path.join(requirement_folder, f"{requirement_name}_EpicDetails.txt")

    # Write the response text to the file
    with open(file_path, "w") as file:
        file.write(response_text)

def split_epic_details1(requirement_folder, epic_chunk_folder, appname):
    epic_file = os.path.join(requirement_folder, f'{appname}_EpicDetails.txt')
    os.makedirs(epic_chunk_folder, exist_ok=True)
    with open(epic_file, 'r') as file:
        content = file.read()

    # Split the content into individual epics
    epics = content.strip().split('Epic ')

    # Process each epic
    for epic in epics:
        if not epic.strip():
            continue

        # Restore 'Epic ' to the beginning of each epic
        if not epic.startswith('Epic ') or epic.startswith('- Epic '):
            epic = 'Epic ' + epic

        # Extract the epic name from the first line of the epic
        first_line = epic.split('\n', 1)[0]
        epic_number, epic_name = first_line.split(':', 1)
        epic_number = epic_number.strip()
        epic_name = epic_name.strip()

        # Create the file name by cleaning up the epic name
        filename = f'{epic_number}_{epic_name.replace(" ", "_")}.txt'
        filepath = os.path.join(epic_chunk_folder, filename)

        # Write the details to the file
        with open(filepath, 'w') as epic_file:
            epic_file.write(epic.strip())

def split_epic_details(requirement_folder, epic_chunk_folder, appname):
    epic_file = os.path.join(requirement_folder, f'{appname}_EpicDetails.txt')
    os.makedirs(epic_chunk_folder, exist_ok=True)
    
    with open(epic_file, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Define regex patterns for detecting Epic and Feature lines
    epic_pattern = re.compile(r'^\s*-?\s*Epic\s*(\d+):?\s*(.*)?$', re.IGNORECASE)
    feature_pattern = re.compile(r'^\s*-?\s*Feature\s*(\d+\.\d+):\s*(.*)$', re.IGNORECASE)
    
    # Split the content into lines
    lines = content.splitlines()
    
    current_epic = None
    epic_details = []
    
    for line in lines:
        epic_match = epic_pattern.match(line)
        if epic_match:
            epic_number = epic_match.group(1).strip()
            epic_name = epic_match.group(2).strip() if epic_match.group(2) else f"Epic {epic_number}"
            current_epic = f"Epic {epic_number}: {epic_name}"
            epic_details.append({"epic": current_epic, "content": []})
            continue
        
        feature_match = feature_pattern.match(line)
        if feature_match and current_epic:
            feature_number = feature_match.group(1).strip()
            feature_name = feature_match.group(2).strip()
            epic_details[-1]["content"].append(f"Feature {feature_number}: {feature_name}")
    
    # Process and save each epic
    for epic_detail in epic_details:
        epic_name = epic_detail["epic"].split(':', 1)[1].strip()
        epic_number = epic_detail["epic"].split(':', 1)[0].split()[1]
        filename = f'{epic_number}_{epic_name.replace(" ", "_")}.txt'
        filepath = os.path.join(epic_chunk_folder, filename)
        
        with open(filepath, 'w', encoding='utf-8') as epic_file:
            epic_file.write(epic_detail["epic"] + '\n')
            for feature in epic_detail["content"]:
                epic_file.write(feature + '\n')

def merge_txt_files_to_word(epic_response_folder,output_word_file):
    # Create a new Word document
    document = Document()
    
    # Iterate through each file in the input folder
    for filename in sorted(os.listdir(epic_response_folder)):
        # Construct the full path of the file
        file_path = os.path.join(epic_response_folder, filename)
        
        # Check if it's a file
        if os.path.isfile(file_path):
            # Add the filename as a heading in the document
            # document.add_heading(filename, level=1)
            
            # Read the content of the file
            with open(file_path, 'r') as file:
                content = file.read()
            
            # Add the content to the document
            document.add_paragraph(content)
            
            # Add a page break after each file
            document.add_page_break()
    
    # Save the document
    # filename = 'Requirement_document.docx'
    if os.path.exists(output_word_file):
        os.remove(output_word_file)

    document.save(output_word_file)
    print(f'Requirement file created')

def extract_requirements_from_txt(requirement_folder,epic_llm_response_folder,usecase_requirement_file):
    # usecase_requirement = os.path.join(requirement_folder, filename)
    feature_file = os.path.join(requirement_folder, 'Feature file.csv')
    if os.path.exists(usecase_requirement_file):
        os.remove(usecase_requirement_file)
        print(f'Deleted existing {usecase_requirement_file}')

    # Initialize variables to hold the extracted data
    epics = []
    epic_name = []
    features = []
    feature_name = []
    tasks = []
    acceptance_criteria = []

    # Process each text file in the epic_llm_response folder
    folder_path = epic_llm_response_folder
    for filename in sorted(os.listdir(folder_path)):
        file_path = os.path.join(folder_path, filename)

        # Read the content of each text file
        with open(file_path, 'r') as file:
            lines = file.readlines()
        
        # Temporary variables to store current epic and feature details
        current_epic_no = ""
        current_epic_name = ""
        current_feature_no = ""
        current_feature_name = ""
        current_task = ""
        current_acceptance_criteria = ""
        capturing_acceptance_criteria = False
        nested_bullet = False
        last_added_feature_no = None

        # Parse the text file line by line
        for line in lines:
            text = line.strip()

            if text.startswith("Epic"):
                epic_parts = text.split(":", 1)
                current_epic_no = epic_parts[0].strip().split()[1]
                current_epic_name = epic_parts[1].strip()
                epics.append(current_epic_no)
                epic_name.append(current_epic_name)

            elif text.startswith("Feature"):
                feature_parts = text.split(":", 1)
                current_feature_no = feature_parts[0].strip().split()[1]
                current_feature_name = feature_parts[1].strip()
                features.append(current_feature_no)
                feature_name.append(current_feature_name)

            elif text.startswith("User story"):
                user_story_parts = text.split(":", 1)
                current_task = user_story_parts[1].strip()
                tasks.append(current_task)
                user_story_id = user_story_parts[0].split()[2]
                if user_story_id.startswith(current_feature_no):
                    features.append(current_feature_no)
                    feature_name.append(current_feature_name)

            elif text.startswith("Acceptance criteria"):
                if capturing_acceptance_criteria and current_acceptance_criteria:
                    acceptance_criteria.append(current_acceptance_criteria.strip())
                capturing_acceptance_criteria = True
                current_acceptance_criteria = ""
                nested_bullet = False

            elif capturing_acceptance_criteria:
                if ":" in text:
                    if nested_bullet:
                        current_acceptance_criteria = current_acceptance_criteria.rstrip(", ") + ". "
                    nested_bullet = True
                    current_acceptance_criteria += text.strip() + " "
                elif nested_bullet:
                    current_acceptance_criteria += text.strip() + ", "
                else:
                    current_acceptance_criteria += text + " "
        
        # Add the last set of acceptance criteria if not already added
        if capturing_acceptance_criteria and current_acceptance_criteria:
            if nested_bullet:
                current_acceptance_criteria = current_acceptance_criteria.rstrip(", ") + "."
            acceptance_criteria.append(current_acceptance_criteria.strip())
    
    print(f'Len of task : {len(tasks)}')
    print(f'Len of ac : {len(acceptance_criteria)}')
    data = {
        'Task': tasks,
        'Acceptance Criteria': acceptance_criteria
    }

    spec_features = {
        'Feature No' : features,
        'Feature Name' : feature_name
    }
    df = pd.DataFrame(data)
    spec_df = pd.DataFrame(spec_features)

    counts = spec_df.groupby(spec_df.columns.tolist()).size()

    # Filter out rows to keep based on their counts
    rows_to_keep = counts[counts > 1].index
    filtered_df = spec_df[spec_df[spec_df.columns.tolist()].apply(tuple, axis=1).isin(rows_to_keep)]

    # Remove one occurrence of each duplicated row
    df_cleaned = spec_df.copy()
    for row in rows_to_keep:
        indices_to_remove = filtered_df[filtered_df[spec_df.columns.tolist()].apply(tuple, axis=1) == row].index
        if len(indices_to_remove) > 1:
            indices_to_remove = indices_to_remove[0]  # Remove only the first occurrence
            df_cleaned = df_cleaned.drop(indices_to_remove)

    epic_dict = dict(zip(epics, epic_name))
    df_cleaned['Epic No'] = df_cleaned['Feature No'].apply(lambda x: x.split('.')[0])
    df_cleaned['Epic Name'] = df_cleaned['Epic No'].apply(lambda x: epic_dict.get(x, ''))
    df_cleaned['Task'] = tasks
    df_cleaned['Acceptance Criteria'] = acceptance_criteria

    df_cleaned = df_cleaned[['Epic No', 'Epic Name', 'Feature No', 'Feature Name', 'Task', 'Acceptance Criteria']]

    df_cleaned.to_csv(usecase_requirement_file, index=False)
    print(f'The {usecase_requirement_file} is:')
    print(df_cleaned)
    print(f'New user requirement captured')
    unique_epic = set(epics)
    unique_epic_name = set(epic_name)

    print(f'Unique epic: {unique_epic}')
    print(f'Unique epic name: {unique_epic_name}')
    
    return unique_epic, unique_epic_name

def transfer_python_file(source_path, destination_path):
    try:
        # Ensure the source file exists and is a Python file
        if not os.path.isfile(source_path) or not source_path.endswith('.py'):
            print("The source file does not exist or is not a Python file.")
            return None

        # Ensure the destination directory exists
        if not os.path.exists(destination_path):
            os.makedirs(destination_path)

        # Move the file
        new_file_path = shutil.move(source_path, destination_path)
        print(f"File moved successfully to {new_file_path}")
        return new_file_path

    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def copy_python_file(source_path, destination_path):
    try:
        # Ensure the source file exists and is a Python file
        if not os.path.isfile(source_path) or not source_path.endswith('.py'):
            print("The source file does not exist or is not a Python file.")
            return None

        # Ensure the destination directory exists
        if not os.path.exists(destination_path):
            os.makedirs(destination_path)

        # Copy the file
        new_file_path = shutil.copy(source_path, destination_path)
        print(f"File copied successfully to {new_file_path}")
        return new_file_path

    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def merge_txt_files(source_folder, destination_folder, output_filename):
    try:
        # Ensure the destination directory exists
        os.makedirs(destination_folder,exist_ok=True)
        # Create the path for the merged file
        merged_file_path = os.path.join(destination_folder, output_filename)

        # Open the output file in write mode
        with open(merged_file_path, 'w', encoding='utf-8') as outfile:
            # Iterate through all files in the source folder
            for filename in os.listdir(source_folder):
                file_path = os.path.join(source_folder, filename)
                
                # Check if the current file is a .txt file
                if os.path.isfile(file_path) and filename.endswith('.txt'):
                    # Read each .txt file and write its contents to the merged file
                    with open(file_path, 'r', encoding='utf-8') as infile:
                        outfile.write(infile.read())
                        outfile.write("\n")  # Optional: Add a newline between files

        print(f"Merged file created at {merged_file_path}")
        return merged_file_path

    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def open_streamlit_app():
    # Assuming you're running this locally; otherwise, use the deployed URL
    new_url = "http://localhost:8507"  # Change this to your Streamlit app URL if hosted
    webbrowser.open_new_tab(new_url)